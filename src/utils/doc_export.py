from __future__ import annotations
from pathlib import Path
from typing import Dict, Iterable
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def _ensure_dir(p: Path) -> None:
    p.parent.mkdir(parents=True, exist_ok=True)

def export_docx(sections: Dict[str, str], out_path: str | Path = "output/docs/IT_STORM_RAG_Report.docx") -> Path:
    """
    Crée un DOCX propre avec titres H1 et contenu par section.
    sections: dict[title] = content (texte ou liste en lignes '- ').
    """
    out = Path(out_path)
    _ensure_dir(out)

    doc = Document()
    # Titre
    title = doc.add_heading("IT STORM — RAG POC Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Style body
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for title, content in sections.items():
        doc.add_heading(title, level=1)
        # Si liste à puces (lignes commençant par "- "), on crée des items
        lines = [l for l in content.splitlines() if l.strip()]
        if lines and all(l.strip().startswith("-") for l in lines):
            for ln in lines:
                p = doc.add_paragraph(ln.lstrip("- ").strip(), style="List Bullet")
                p_format = p.paragraph_format
                p_format.space_after = Pt(4)
        else:
            # Paragraphe normal (on split par doubles sauts de ligne pour aérer)
            blocks: Iterable[str] = "\n".join(lines).split("\n\n") if lines else [content]
            for block in blocks:
                p = doc.add_paragraph(block.strip())
                p_format = p.paragraph_format
                p_format.space_after = Pt(8)

    doc.save(str(out))
    return out
