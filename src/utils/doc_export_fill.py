# src/ui/utils/doc_export_fill.py
from __future__ import annotations

import os
from pathlib import Path
from datetime import datetime
from typing import Dict, Optional

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception as e:
    raise RuntimeError(
        "Le module 'python-docx' est requis. "
        "Installe-le avec: pip install python-docx"
    ) from e


def _slugify(name: str) -> str:
    keep = "-_(). "
    slug = "".join(
        c if c.isalnum() or c in keep else "-"
        for c in (name or "").strip()
    )
    # compact multiple dashes
    while "--" in slug:
        slug = slug.replace("--", "-")
    return slug.strip().strip("-")


def export_docx_fill(
    sections: Dict[str, str],
    *,
    title: str = "IT_STORM_Proposal_EN",
    client: Optional[str] = None,
    output_path: Optional[str] = None,
    output_dir: str = "outputs",
) -> str:
    """
    Exporte un document DOCX avec les sections auto-remplies du Tab 3.

    Args:
        sections: dict { "Section Title": "Content ..." }
        title: titre du document (champ 'Document title' du Tab 3)
        client: nom du client (champ optionnel du Tab 3)
        output_path: chemin complet du fichier DOCX à créer (prioritaire si fourni)
        output_dir: dossier de sortie par défaut si output_path n'est pas fourni

    Returns:
        str: chemin absolu du fichier DOCX créé.
    """
    # Déterminer le chemin de sortie
    if output_path:
        out_path = Path(output_path)
        out_dir = out_path.parent
    else:
        out_dir = Path(output_dir)
        out_dir.mkdir(parents=True, exist_ok=True)
        filename = f"{_slugify(title or 'IT_STORM_Proposal_EN')}.docx"
        out_path = out_dir / filename

    # Créer le document
    doc = Document()

    # Titre centré
    heading = doc.add_heading(title or "IT_STORM_Proposal_EN", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Sous-titre client / date
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta_p.add_run(
        f"{'Client: ' + client + ' — ' if client else ''}"
        f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    run.italic = True
    run.font.size = Pt(10)

    # Lignes vides
    doc.add_paragraph("")

    # Contenu des sections
    for sec_title, content in (sections or {}).items():
        if not sec_title:
            continue
        doc.add_heading(str(sec_title), level=1)
        paragraph = doc.add_paragraph(content.strip() if content else "")
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_after = Pt(6)

    # Pied de page simple
    try:
        section = doc.sections[0]
        footer = section.footer.paragraphs[0]
        footer.text = "StormCopilot — IT STORM · Auto-generated (RAG EN)"
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        # Footer best-effort; on ignore si l'API n'est pas dispo
        pass

    # Sauvegarde
    doc.save(str(out_path))

    # Retourne un chemin absolu propre (Windows/Unix)
    return str(out_path.resolve())
