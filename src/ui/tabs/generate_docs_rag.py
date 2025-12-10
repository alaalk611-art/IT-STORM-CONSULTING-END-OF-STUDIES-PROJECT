# -*- coding: utf-8 -*-
# ============================================================
# Path: src/ui/tabs/generate_docs_rag.py
# Role: Streamlit tab – RAG ancré + Jury multi-LLM via Ollama
# + Résumé de documents aligné sur src/rag_sum.py
#   - Question → smart_rag_answer / ask_multi_ollama (comme avant)
#   - Résumé → même logique que la CLI: _summarize_text_three_cli
# ============================================================

from __future__ import annotations
import io
import re
import os
import time
import math
from collections import Counter
from typing import List, Dict, Any, Optional, Tuple
from docx import Document
import streamlit as st
import pandas as pd

# --- RAG strict engine (imports robustes) ---
try:
    import src.rag_brain as _rb
    smart_rag_answer = getattr(_rb, "smart_rag_answer", None)
    ask_multi_ollama = getattr(_rb, "ask_multi_ollama", None)
except Exception:
    smart_rag_answer = None
    ask_multi_ollama = None

# --- Résumé (module dédié) ---
try:
    # on importe AUSSI read_any_text + _summarize_text_three_cli pour coller au mode CLI
    from src.rag_sum import (
        summarize_file,
        summary_quality_report,
        read_any_text,
        _summarize_text_three_cli,
    )
except Exception:
    summarize_file = None
    summary_quality_report = None
    read_any_text = None
    _summarize_text_three_cli = None

# =================== CSS GLOBAL ===================
st.markdown("""
<style>
/* ===== DataFrame : montrer TOUTE la réponse sans coupe ===== */
[data-testid="stDataFrame"] div[role="gridcell"],
[data-testid="stDataFrame"] div[role="gridcell"] p{
  white-space: pre-wrap !important;
  overflow: visible !important;
  text-overflow: initial !important;
  line-height: 1.35 !important;
  word-break: break-word !important;
  line-break: anywhere !important;
}

/* ===== Helper “?” cliquable (tooltip) ===== */
.helper-row{
  display:flex; align-items:center; justify-content:space-between; gap:.5rem;
  margin:.15rem 0 .35rem 0;
}
.helper-label{ font-weight:600; color:#0f172a; }
.helper{
  position:relative; display:inline-flex; align-items:center; justify-content:center;
  width:22px; height:22px; border-radius:50%;
  background:#e5eefc; color:#284b9f; font-weight:700; cursor:pointer;
  border:1px solid #c7d2fe; user-select:none;
}
.helper:hover{ background:#dbe8ff; }
.helper .tooltip{
  position:absolute; top:30px; right:-8px;
  min-width:220px; max-width:360px; padding:.55rem .7rem; border-radius:.55rem;
  border:1px solid #e5e7eb; background:#fff; color:#111827;
  box-shadow:0 8px 18px rgba(0,0,0,.08); font-size:.86rem; line-height:1.25;
  display:none; z-index:20;
}
.helper:focus .tooltip, .helper:hover .tooltip{ display:block; }
</style>
""", unsafe_allow_html=True)

# ============================================================
# Métadonnées modèles (affichage léger)
# ============================================================

_MODEL_META: Dict[str, Dict[str, str]] = {
    "mistral:7b-instruct": {"emoji": "🦅", "color": "#2b7cff"},
    "llama3.2:3b":         {"emoji": "🦙", "color": "#7b5cff"},
    "qwen2.5:7b":          {"emoji": "🐉", "color": "#16a34a"},
    "tinyllama:latest":    {"emoji": "🌱", "color": "#64748b"},
    "smart_rag_answer":    {"emoji": "🔎", "color": "#334155"},
}

def _default_models() -> List[str]:
    return ["mistral:7b-instruct", "llama3.2:3b", "qwen2.5:7b", "tinyllama:latest"]

def _label(name: str, rank: Optional[int] = None) -> str:
    meta = _MODEL_META.get(name, {"emoji": "🤖", "color": "#334155"})
    emo = meta["emoji"]
    badge = f" · Rang {rank}" if rank else ""
    return f"{emo} {name}{badge}"

def _decorate_model_label(name: str, rank: Optional[int] = None) -> str:
    meta = _MODEL_META.get(name, {"emoji": "🤖", "color": "#334155"})
    emoji = meta["emoji"]; color = meta["color"]
    medal = ""
    if rank is not None:
        if rank == 1: medal = "🥇"
        elif rank == 2: medal = "🥈"
        elif rank == 3: medal = "🥉"
        else: medal = f"🏁{rank}"
    return f"""
<span style="display:inline-flex;align-items:center;gap:.55rem;padding:.35rem .65rem;
      border:1px solid #e5e7eb;background:#f8fafc;border-radius:999px;font-weight:600;">
  <span style="width:12px;height:12px;border-radius:50%;display:inline-block;background:{color};"></span>
  <span>{emoji} {name}</span>
  {"<span>"+medal+"</span>" if medal else ""}
</span>
""".strip()

# ============================================================
# Helpers de base (priorité: contenu)
# ============================================================

def _safe_num(x, default: float = 0.0) -> float:
    try:
        return float(x)
    except Exception:
        return default

def _score_from_result(r: Dict[str, Any]) -> float:
    if "score" in r:
        return _safe_num(r["score"])
    if "confidence" in r:
        return _safe_num(r["confidence"])
    return 0.0

def _legend_metrics() -> None:
    lang = st.session_state.get("lang", "fr")
    c1, c2, c3, c4 = st.columns(4)
    if lang == "en":
        with c1: st.markdown("**Confidence**", help="Global reliability of the generated answer (0–1).")
        with c2: st.markdown("**Grounding**", help="How well the answer is anchored in the retrieved context (0–1).")
        with c3: st.markdown("**Coverage**", help="Diversity and completeness of the answer versus the retrieved context (0–1).")
        with c4: st.markdown("**Style**", help="Readability and formatting quality of the generated text (0–1).")
    else:
        with c1: st.markdown("**Confiance**", help="Fiabilité globale de la réponse générée (0–1).")
        with c2: st.markdown("**Grounding**", help="Degré d’ancrage de la réponse sur les extraits sources (0–1).")
        with c3: st.markdown("**Couverture**", help="Diversité et exhaustivité du contenu par rapport au contexte (0–1).")
        with c4: st.markdown("**Style**", help="Lisibilité et clarté de la rédaction (0–1).")

# ============================================================
# Slider Top-K dans la sidebar (avec helper “?”)
# ============================================================
def _make_docx_from_summary(title: str, intro: str, body: str) -> io.BytesIO:
    """
    Crée un fichier DOCX IT-STORM avec :
      - Page de garde (cover page) : logo + titre + sous-titre
      - Deuxième page : résumé détaillé avec header/footer corporate
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_SECTION_START

    # ==============================
    # DOCUMENT & COULEURS
    # ==============================
    doc = Document()

    BLUE = RGBColor(37, 99, 235)
    DARK = RGBColor(55, 65, 81)
    LIGHT = RGBColor(107, 114, 128)
    BORDER = RGBColor(209, 213, 219)

    # Style global
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Calibri"
    style_normal.font.size = Pt(11)

    # ==============================
    # SECTION 1 : COVER PAGE
    # ==============================
    cover_section = doc.sections[0]
    cover_section.top_margin = Inches(1.5)
    cover_section.bottom_margin = Inches(1.0)
    cover_section.left_margin = Inches(1.2)
    cover_section.right_margin = Inches(1.2)

    # Logo centré
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_logo = p_logo.add_run()
    try:
        # Adapte ce chemin à ton projet
        run_logo.add_picture("assets/itstorm_logo.png", width=Inches(1.5))
    except Exception:
        pass

    # Espace après logo
    doc.add_paragraph("")

    # Titre principal
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run(title or "Résumé de document")
    r_title.bold = True
    r_title.font.size = Pt(28)
    r_title.font.color.rgb = BLUE

    # Sous-titre (intro)
    if intro:
        p_intro = doc.add_paragraph()
        p_intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_intro = p_intro.add_run(intro.strip())
        r_intro.font.size = Pt(12)
        r_intro.font.color.rgb = DARK
        r_intro.italic = True

    # Ligne bleue décorative
    doc.add_paragraph("")
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_line = p_line.add_run("────────────────────────────────────────────")
    r_line.font.color.rgb = BLUE

    # Baseline IT-STORM en bas de la cover
    doc.add_paragraph("")
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("IT-STORM · Intelligent Consulting Copilot")
    r_sub.font.size = Pt(11)
    r_sub.font.color.rgb = LIGHT

    # Petite marge verticale
    doc.add_paragraph("")

    # Page suivante (section 2)
    main_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    main_section.top_margin = Inches(0.8)
    main_section.bottom_margin = Inches(0.8)
    main_section.left_margin = Inches(0.9)
    main_section.right_margin = Inches(0.9)

    # ==============================
    # SECTION 2 : HEADER & FOOTER
    # ==============================
    header = main_section.header
    if header.paragraphs:
        h = header.paragraphs[0]
        for run in h.runs:
            run.text = ""
    else:
        h = header.add_paragraph()

    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h_run_logo = h.add_run()
    try:
        h_run_logo.add_picture("assets/itstorm_logo.png", width=Inches(0.9))
    except Exception:
        pass

    h.add_run("  ")
    h_title = h.add_run("IT-STORM")
    h_title.bold = True
    h_title.font.size = Pt(12)
    h_title.font.color.rgb = BLUE

    h_sub = h.add_run("  ·  Intelligent Consulting Copilot")
    h_sub.font.size = Pt(9)
    h_sub.font.color.rgb = DARK

    # Footer
    footer = main_section.footer
    if footer.paragraphs:
        f = footer.paragraphs[0]
        for run in f.runs:
            run.text = ""
    else:
        f = footer.add_paragraph()

    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = f.add_run("IT-STORM · Intelligent Consulting Copilot · Document de résumé généré automatiquement.")
    rf.italic = True
    rf.font.size = Pt(8)
    rf.font.color.rgb = LIGHT

    # ==============================
    # SECTION 2 : CONTENU "RÉSUMÉ DÉTAILLÉ"
    # (tout ce qui suit est automatiquement dans la nouvelle section)
    # ==============================

    # Petite ligne bleu clair en haut de la page 2
    p_top_line = doc.add_paragraph()
    r_top = p_top_line.add_run("────────────────────────────────────────────")
    r_top.font.color.rgb = BORDER

    doc.add_paragraph("")

    # Titre de section
    p_sec_title = doc.add_paragraph()
    p_sec_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_sec_title = p_sec_title.add_run("Résumé détaillé")
    r_sec_title.bold = True
    r_sec_title.font.size = Pt(14)
    r_sec_title.font.color.rgb = BLUE

    # Ligne fine grise
    p_sec_line = doc.add_paragraph()
    r_sec_line = p_sec_line.add_run("────────────────────────────")
    r_sec_line.font.color.rgb = BORDER

    doc.add_paragraph("")

    # Corps du résumé : un paragraphe par bloc
    for block in (body or "").split("\n\n"):
        block = block.strip()
        if not block:
            continue
        p = doc.add_paragraph(block)
        p.style = doc.styles["Normal"]
        p_format = p.paragraph_format
        p_format.space_after = Pt(8)
        p_format.line_spacing = 1.35

    # ==============================
    # EXPORT
    # ==============================
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def _sidebar_topk_slider(default:int=6) -> int:
    lang = st.session_state.get("lang", "en")
    is_en = (lang == "en")
    label_txt = "Top-K results" if is_en else "Top-K résultats"
    help_txt  = (
        "Number of most similar document chunks retrieved from the vector database to build the context for each question."
        if is_en else
        "Nombre de segments de documents similaires utilisés pour construire le contexte des réponses."
    )
    with st.sidebar:
        st.markdown(
            f"""
<div class="helper-row">
  <div class="helper-label">{label_txt}</div>
  <div class="helper" tabindex="0">?
    <div class="tooltip">{help_txt}</div>
  </div>
</div>
""",
            unsafe_allow_html=True
        )
        value = st.slider(" ", min_value=2, max_value=12, value=default, step=1, key="topk_sidebar_slider")
    return int(value)

# ============================================================
# Appels moteur
# ============================================================

def _call_smart_rag(question: str, lang: str = "fr", mode: str = "definition") -> Dict[str, Any]:
    if not (smart_rag_answer and callable(smart_rag_answer)):
        return {}
    try:
        return smart_rag_answer(question=question, lang=lang, mode=mode) or {}
    except Exception as e:
        st.error(f"smart_rag_answer a échoué: {e}")
        return {}

def _call_jury(question: str, models: List[str], topk_context: int = 6, timeout: float = 600.0) -> Dict[str, Any]:
    if not (ask_multi_ollama and callable(ask_multi_ollama)):
        return {}
    try:
        return ask_multi_ollama(question=question, models=models,
                                topk_context=int(topk_context), timeout=float(timeout)) or {}
    except TypeError:
        try:
            return ask_multi_ollama(question=question, models=models) or {}
        except Exception as e:
            st.warning(f"Jury indisponible: {e}")
            return {}
    except Exception as e:
        st.warning(f"Jury indisponible: {e}")
        return {}

def _ask_one_model(question: str, model: str, topk_context: int = 6, timeout: float = 600.0) -> Dict[str, Any]:
    jury = _call_jury(question, models=[model], topk_context=topk_context, timeout=timeout)
    res = (jury.get("results") or [])
    if res:
        r = res[0]
        out = {
            "model": r.get("model", model),
            "answer": r.get("answer", ""),
            "score": _safe_num(r.get("score", 0.0)),
            "confidence": _safe_num((r.get("metrics") or {}).get("confidence", r.get("score", 0.0))),
            "grounding": _safe_num((r.get("metrics") or {}).get("grounding", 0.0)),
            "coverage": _safe_num((r.get("metrics") or {}).get("coverage", 0.0)),
            "style": _safe_num((r.get("metrics") or {}).get("style", 0.0)),
            "time": _safe_num(r.get("time", 0.0)),
        }
        return out

    base = _call_smart_rag(question) or {}
    qual = base.get("quality") or {}
    return {
        "model": model,
        "answer": base.get("answer", ""),
        "score": _safe_num(base.get("confidence", 0.0)),
        "confidence": _safe_num(base.get("confidence", 0.0)),
        "grounding": _safe_num(qual.get("grounding", 0.0)),
        "coverage": _safe_num(qual.get("coverage", 0.0)),
        "style": _safe_num(qual.get("style", 0.0)),
        "time": 0.0,
    }

# ============================================================
# Lecture upload de secours (TXT/PDF) – rarement utilisé
# ============================================================

def _read_uploaded_file(upload) -> Tuple[str, str]:
    if not upload:
        return "", ""
    name = (upload.name or "").lower()
    if name.endswith(".txt"):
        try:
            raw = upload.read()
            return raw.decode("utf-8", errors="ignore"), ".txt"
        except Exception:
            return "", ".txt"
    if name.endswith(".pdf"):
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(upload.read()))
            pages = []
            for p in reader.pages:
                try:
                    pages.append(p.extract_text() or "")
                except Exception:
                    pages.append("")
            return "\n".join(pages), ".pdf"
        except Exception:
            return "", ".pdf"
    return "", ""

# ============================================================
# Helpers qualité locaux (fallback si besoin)
# ============================================================

def _tok(s: str) -> list[str]:
    return re.findall(r"\w+", (s or "").lower(), flags=re.UNICODE)

def _ngrams(tokens: list[str], n: int) -> list[tuple]:
    return [tuple(tokens[i:i+n]) for i in range(len(tokens)-n+1)] if n > 0 else []

def _rouge_prf(ref: str, hyp: str, n: int = 1) -> dict:
    r_toks, h_toks = _tok(ref), _tok(hyp)
    r_ngr, h_ngr = Counter(_ngrams(r_toks, n)), Counter(_ngrams(h_toks, n))
    overlap = sum((r_ngr & h_ngr).values())
    R = overlap / (sum(r_ngr.values()) or 1)
    P = overlap / (sum(h_ngr.values()) or 1)
    F = (2*P*R)/(P+R) if (P+R) > 0 else 0.0
    return {"P": P, "R": R, "F1": F}

def _keyword_overlap(ref: str, hyp: str, k: int = 20) -> float:
    rt = [t for t in _tok(ref) if len(t) > 2]
    ht = set([t for t in _tok(hyp) if len(t) > 2])
    if not rt or not ht:
        return 0.0
    top = [w for w, _ in Counter(rt).most_common(k)]
    hit = sum(1 for w in top if w in ht)
    return hit / max(1, len(top))

def _cosine_sim(ref: str, hyp: str) -> float:
    r, h = Counter(_tok(ref)), Counter(_tok(hyp))
    if not r or not h:
        return 0.0
    inter = set(r) & set(h)
    num = sum(r[w]*h[w] for w in inter)
    nr = math.sqrt(sum(v*v for v in r.values()))
    nh = math.sqrt(sum(v*v for v in h.values()))
    return num / (nr*nh) if nr*nh > 0 else 0.0

def _unsupported_sentences(ref: str, hyp: str, thr: float = 0.15) -> list[str]:
    ref_3 = set(_ngrams(_tok(ref), 3))
    out = []
    for sent in re.split(r"(?<=[\.\!\?])\s+", (hyp or "").strip()):
        if not sent:
            continue
        s3 = _ngrams(_tok(sent), 3)
        if not s3:
            continue
        overlap = sum(1 for g in s3 if g in ref_3) / len(s3)
        if overlap < thr:
            out.append(sent.strip())
    return out

def _summary_quality_report(source_text: str, summary_text: str) -> dict:
    Ls, Lh = len(_tok(source_text)), len(_tok(summary_text))
    comp = (1 - (Lh / max(1, Ls))) if Ls else 0.0
    rouge1 = _rouge_prf(source_text, summary_text, n=1)
    rouge2 = _rouge_prf(source_text, summary_text, n=2)
    sim   = _cosine_sim(source_text, summary_text)
    kw    = _keyword_overlap(source_text, summary_text, k=20)
    unsup = _unsupported_sentences(source_text, summary_text, thr=0.15)
    return {
        "tokens_source": Ls,
        "tokens_summary": Lh,
        "compression": comp,
        "cosine": sim,
        "rouge1_f1": rouge1["F1"],
        "rouge2_f1": rouge2["F1"],
        "keyword_overlap": kw,
        "unsupported": unsup,
    }

# ============================================================
# Cohérence de paragraphe (GÉNÉRIQUE, GLOBAL)
# ============================================================

def _cohere_paragraph(raw: str) -> str:
    """
    Nettoyage GÉNÉRIQUE du résumé, sans ajouter de contenu nouveau :
    - enlève les puces / tirets
    - remet les phrases à la ligne en un seul paragraphe
    - capitalise la première lettre
    - ajoute un point final si nécessaire

    Aucun mot ni domaine n'est injecté (fonction 100 % globale).
    """
    if not raw:
        return ""

    lines: List[str] = []
    for ln in raw.splitlines():
        ln = ln.strip(" •-*—\t")
        if ln:
            lines.append(ln.strip())

    if not lines:
        return raw.strip()

    sentences: List[str] = []
    for s in lines:
        s = s.strip()
        if not s:
            continue
        if s[0].isalpha():
            s = s[0].upper() + s[1:]
        if s[-1] not in ".!?…":
            s += "."
        sentences.append(s)

    paragraph = " ".join(sentences)
    return " ".join(paragraph.split())

def _apply_coherence(summary_text: str, use_template: bool) -> str:
    """
    - Si use_template = False → on retourne le texte brut du modèle (juste strip).
    - Si use_template = True  → on applique _cohere_paragraph (global, neutre).
    """
    summary_text = (summary_text or "").strip()
    if not summary_text:
        return ""
    if not use_template:
        return summary_text
    return _cohere_paragraph(summary_text)

# ============================================================
# Fusion LLM des résumés (Ultimate Version)
# ============================================================

def _build_ultimate_summary(
    source_text: str,
    candidates: List[Dict[str, Any]],
    use_template: bool,
    fusion_model: str = "qwen2.5:7b",
    timeout: float = 240.0,
) -> str:
    """
    Construit une *version ultime* en fusionnant les meilleurs résumés de modèles.

    Logique PRO :
      - On envoie au LLM :
          * un extrait (snippet) du texte source
          * tous les résumés candidats (triés par score)
      - On demande un résumé unique, plus riche que chaque résumé isolé,
        en intégrant les informations importantes propres à chaque modèle,
        sans rien inventer.

      - Si l'appel LLM échoue OU si le résultat est quasi identique
        au meilleur résumé, on applique un fallback déterministe :
          * base = meilleur résumé
          * + phrases supplémentaires uniques (non présentes dans la base)
            prélevées dans les autres résumés.

    Aucun contenu extérieur n’est injecté : on ne fait que réordonner /
    combiner des phrases déjà présentes dans les résumés.
    """

    # --------- Helpers internes : split / fusion déterministe ---------
    def _split_sentences(txt: str) -> List[str]:
        txt = (txt or "").strip()
        if not txt:
            return []
        # split grossier en phrases
        parts = re.split(r"(?<=[\.!\?…])\s+", txt)
        out = []
        for s in parts:
            s = s.strip()
            if not s:
                continue
            out.append(s)
        return out

    def _simple_deterministic_fusion(cands: List[Dict[str, Any]], max_chars: int = 2200) -> str:
        """
        Fusion "safe" sans LLM :
          - on prend le meilleur résumé comme base
          - on ajoute des phrases uniques venant des autres résumés
          - aucun mot inventé (tout vient des candidats)
        """
        if not cands:
            return ""

        # Meilleur candidat par score
        best_loc = max(cands, key=lambda x: x.get("score", 0.0))
        base = (best_loc.get("answer", "") or "").strip()
        if not base:
            return ""

        base_sents = _split_sentences(base)
        merged_sents: List[str] = list(base_sents)
        merged_text = " ".join(merged_sents)

        # On parcourt les autres résumés, du meilleur au moins bon
        for c in sorted(cands, key=lambda x: x.get("score", 0.0), reverse=True):
            if c is best_loc:
                continue
            ans = (c.get("answer", "") or "").strip()
            if not ans:
                continue
            for s in _split_sentences(ans):
                s_clean = s.strip()
                # on ignore les phrases trop courtes ou déjà présentes
                if len(s_clean) < 40:
                    continue
                if s_clean in merged_text:
                    continue
                merged_sents.append(s_clean)
                merged_text = " ".join(merged_sents)
                if len(merged_text) >= max_chars:
                    break
            if len(merged_text) >= max_chars:
                break

        return merged_text.strip()

    # --------- Normalisation candidats ---------
    source_text = (source_text or "").strip()
    candidates = [c for c in (candidates or []) if (c.get("answer") or "").strip()]

    if not candidates:
        return ""

    # Si un seul résumé → rien à fusionner
    if len(candidates) == 1:
        raw = (candidates[0].get("answer", "") or "").strip()
        return _apply_coherence(raw, use_template)

    # --------- Snippet du texte source (générique) ----------
    if source_text:
        n = len(source_text)
        if n <= 8000:
            limit = 8000
        elif n <= 20000:
            limit = 12000
        else:
            limit = 18000

        snippet = source_text[:limit]
        if len(source_text) > limit:
            snippet += "\n[Texte source tronqué pour la fusion]"
    else:
        snippet = ""

    # --------- Bloc des résumés candidats ----------
    # Triés par score pour mettre les meilleurs en premier
    bullets: List[str] = []
    for idx, c in enumerate(sorted(candidates, key=lambda x: x.get("score", 0.0), reverse=True), start=1):
        model_name = c.get("model", "—")
        ans = (c.get("answer", "") or "").strip()
        if not ans:
            continue
        bullets.append(f"### Résumé {idx} ({model_name})\n{ans}")

    if not bullets:
        # Fallback brutal : meilleur résumé
        best_loc = max(candidates, key=lambda x: x.get("score", 0.0))
        raw = (best_loc.get("answer", "") or "").strip()
        return _apply_coherence(raw, use_template)

    bullets_text = "\n\n".join(bullets)

    # --------- Essai 1 : fusion via LLM (Ollama) ----------
    fused_raw = ""
    try:
        try:
            from src.llm.ollama_client import generate_ollama
        except Exception:
            generate_ollama = None

        if generate_ollama is not None:
            prompt = (
                "On te donne plusieurs résumés d'un même document, rédigés par des modèles différents.\n"
                "Ton rôle est de produire un *résumé ultime* en français, fidèle et plus riche que chaque résumé pris isolément.\n\n"
                "Règles à respecter strictement :\n"
                "1) Utilise uniquement les informations présentes dans le texte source et les résumés fournis.\n"
                "   - N'invente aucun fait, chiffre ou concept.\n"
                "2) Intègre les points importants propres à chaque résumé, même s'ils ne sont présents que dans un seul.\n"
                "3) Évite les répétitions : fusionne les idées proches en une formulation claire.\n"
                "4) Style : texte continu, fluide, professionnel, sans listes à puces.\n"
                "5) Structure : 2 à 3 paragraphes maximum.\n"
                "6) Longueur : légèrement plus longue et plus complète que chaque résumé isolé, mais reste concise.\n"
                "7) Ne mentionne jamais les noms des modèles (Mistral, Llama, Qwen, etc.).\n\n"
                "EXTRAIT DU TEXTE SOURCE (éventuellement tronqué) :\n"
                "------------------------------\n"
                f"{snippet}\n"
                "------------------------------\n\n"
                "RÉSUMÉS CANDIDATS :\n"
                f"{bullets_text}\n\n"
                "Maintenant, écris un *résumé ultime* unique et cohérent en français,\n"
                "en respectant strictement les règles ci-dessus.\n"
                "RÉSUMÉ ULTIME :"
            )

            fused_raw = generate_ollama(
                fusion_model,
                prompt,
                temperature=0.15,
                max_tokens=750,
                stream=True,
                timeout=timeout,
                options={
                    "num_ctx": 3072,
                    "top_k": 40,
                    "top_p": 0.9,
                    "repeat_penalty": 1.05,
                    "seed": 1,
                },
            )
            fused_raw = (fused_raw or "").strip()
    except Exception:
        fused_raw = ""

    # --------- Fallback : si LLM KO ou résultat trop proche du best ---------
    best_loc = max(candidates, key=lambda x: x.get("score", 0.0))
    best_text = (best_loc.get("answer", "") or "").strip()

    use_fallback = False
    if not fused_raw:
        use_fallback = True
    else:
        # Très simple heuristique de similarité : égalité exacte
        # ou différence de longueur minime + inclusion
        bt = best_text.replace("\n", " ").strip()
        fr = fused_raw.replace("\n", " ").strip()
        if bt == fr:
            use_fallback = True
        else:
            len_bt = len(bt)
            len_fr = len(fr)
            if abs(len_bt - len_fr) < 40 and (bt in fr or fr in bt):
                use_fallback = True

    if use_fallback:
        merged = _simple_deterministic_fusion(candidates, max_chars=2200)
        merged = merged or best_text
        return _apply_coherence(merged, use_template)

    # --------- Chemin normal : résultat LLM + post-traitement cohérence ---------
    fused = _apply_coherence(fused_raw, use_template)
    return fused if fused else best_text

# ============================================================
# UI (modèles cliquables) + Podium/Course + Tableaux
# ============================================================

def _pick_models() -> List[str]:
    st.caption("Clique sur un modèle pour l’activer ou le désactiver.")
    defaults = _default_models()
    if "selected_models" not in st.session_state:
        st.session_state["selected_models"] = defaults[:3]
    selected = list(st.session_state["selected_models"])

    st.markdown("""
<style>
.model-grid { display:flex; flex-wrap:wrap; gap:10px; margin-bottom:12px; }
.model-card {
  flex: 1 1 calc(25% - 10px);
  border: 2px solid transparent; border-radius: 14px; padding: 10px 14px;
  background: #ffffff; box-shadow: 0 4px 12px rgba(0,0,0,0.05);
  cursor: pointer; display: flex; align-items: center; justify-content:center; gap:8px;
  font-weight: 700; color: #111827; text-align:center; transition: all .2s ease;
}
.model-card:hover { transform: scale(1.03); box-shadow: 0 6px 16px rgba(0,0,0,0.08); }
.model-card.active { background: linear-gradient(135deg,#dbeafe,#bfdbfe); border-color:#2563eb; color:#1e3a8a; }
.model-dot { display:inline-block; width:10px; height:10px; border-radius:50%; }
</style>
""", unsafe_allow_html=True)

    st.markdown('<div class="model-grid">', unsafe_allow_html=True)

    for i, m in enumerate(defaults):
        meta = _MODEL_META.get(m, {"emoji": "🤖", "color": "#334155"})
        active = m in selected
        card_key = f"model_card_{i}"
        color = meta["color"]
        bg_class = "active" if active else ""

        st.markdown(
            f"""
<div class="model-card {bg_class}" id="{card_key}">
  <span class="model-dot" style="background:{color};"></span>
  {meta['emoji']} {m}
</div>
<script>
const el_{i} = window.parent.document.getElementById("{card_key}");
if (el_{i}) {{
  el_{i}.onclick = function() {{
    const api = window.parent.streamlitApi;
    if (api) {{
      api.setComponentValue("{card_key}", !{str(active).lower()});
    }}
  }}
}}
</script>
""",
            unsafe_allow_html=True,
        )

        clicked = st.session_state.get(card_key, None)
        if clicked is not None:
            if active:
                if m in selected:
                    selected.remove(m)
            else:
                if m not in selected:
                    selected.append(m)
            selected = [x for x in _default_models() if x in set(selected)]
            st.session_state["selected_models"] = selected
            st.session_state[card_key] = None
            st.rerun()

    st.markdown('</div>', unsafe_allow_html=True)

    if not selected:
        selected = ["mistral:7b-instruct", "llama3.2:3b", "qwen2.5:7b"]
        st.session_state["selected_models"] = selected

    return selected

def _display_grounded_block(block: Dict[str, Any]) -> None:
    st.markdown("### 🔍 Réponse ancrée")
    ans = block.get("answer") or "Je ne sais pas."
    st.write(ans)

    conf = _safe_num(block.get("confidence", 0.0))
    qual = block.get("quality") or {}
    grd  = _safe_num(qual.get("grounding", 0.0))
    cov  = _safe_num(qual.get("coverage", 0.0))
    sty  = _safe_num(qual.get("style", 0.0))

    _legend_metrics()
    c1, c2, c3, c4 = st.columns(4)
    with c1: st.metric("Confiance", f"{conf:0.2f}")
    with c2: st.metric("Grounding", f"{grd:0.2f}")
    with c3: st.metric("Couverture", f"{cov:0.2f}")
    with c4: st.metric("Style", f"{sty:0.2f}")

    quotes = block.get("quotes") or []
    if quotes:
        with st.expander("Extraits cités"):
            for q in quotes:
                st.write(f"- « {q.get('quote','').strip()} » ({q.get('source','')})")

def _render_podium(results: List[Dict[str, Any]]) -> None:
    if not results:
        return
    st.subheader("🏆 Podium (Top 3)")
    top3 = sorted(results, key=_score_from_result, reverse=True)[:3]
    data = []
    for i, r in enumerate(top3, start=1):
        data.append({"rank": i, "model": r.get("model", "—"), "score": f"{_score_from_result(r):0.3f}"})
    while len(data) < 3:
        data.append({"rank": len(data)+1, "model": "—", "score": "—"})

    html = f"""
<style>
.podium-wrap {{ width:100%; margin:16px 0 18px 0; }}
.podium {{ display:flex; align-items:flex-end; justify-content:center; gap:36px;
          width:100%; height:260px; padding-bottom:10px; }}
.podium .col {{ display:flex; flex-direction:column; align-items:center; gap:14px; }}
.podium .block {{ width:180px; border-radius:14px 14px 0 0;
  background: linear-gradient(180deg,#ffe27a 0%,#f5b400 45%,#b67b00 100%);
  box-shadow:0 6px 16px rgba(0,0,0,.15), inset 0 2px 0 rgba(255,255,255,.4);
  border:1px solid rgba(0,0,0,.08);
}}
.podium .rank1 .block {{ height:160px; background: linear-gradient(180deg,#fff176 0%,#fdd835 45%,#f9a825 100%); }}
.podium .rank2 .block {{ height:130px; background: linear-gradient(180deg,#d1d5db 0%,#9ca3af 45%,#6b7280 100%); }}
.podium .rank3 .block {{ height:110px; background: linear-gradient(180deg,#fbc8a2 0%,#f59e0b 45%,#b45309 100%); }}
.podium .model {{ font-weight:800; font-size:1.05rem; color:#111827; text-align:center; letter-spacing:.3px; }}
.podium .score {{ font-size:1rem; color:#1e3a8a; background:#eef2ff;
  border-radius:999px; padding:6px 16px; font-weight:700; box-shadow:inset 0 1px 2px rgba(0,0,0,.08); }}
.podium .medal {{ font-size:2rem; margin-bottom:-6px; }}
@media (max-width: 900px) {{
  .podium .block {{ width:120px; }}
  .podium .score {{ font-size:0.9rem; padding:4px 10px; }}
  .podium .model {{ font-size:0.95rem; }}
  .podium .medal {{ font-size:1.6rem; }}
}}
</style>

<div class="podium-wrap">
  <div class="podium">
    <div class="col rank2">
      <div class="medal">🥈</div>
      <div class="block"></div>
      <div class="model">{data[1]['model']}</div>
      <div class="score">Score : <b>{data[1]['score']}</b></div>
    </div>
    <div class="col rank1">
      <div class="medal">🥇</div>
      <div class="block"></div>
      <div class="model">{data[0]['model']}</div>
      <div class="score">Score : <b>{data[0]['score']}</b></div>
    </div>
    <div class="col rank3">
      <div class="medal">🥉</div>
      <div class="block"></div>
      <div class="model">{data[2]['model']}</div>
      <div class="score">Score : <b>{data[2]['score']}</b></div>
    </div>
  </div>
</div>
"""
    st.markdown(html, unsafe_allow_html=True)

def _render_race(results: List[Dict[str, Any]], steps: int = 12, sleep: float = 0.06) -> None:
    if not results or len(results) < 3:
        return
    st.subheader("🏁 Course (visualisation rapide)")
    rs = sorted(results, key=_score_from_result, reverse=True)[:5]
    max_score = max(_score_from_result(r) for r in rs) or 1.0
    cols = st.columns(len(rs))
    bars = []
    for i, (c, r) in enumerate(zip(cols, rs), start=1):
        with c:
            st.markdown(_decorate_model_label(r.get("model", "—"), rank=i if i <= 3 else None), unsafe_allow_html=True)
            bars.append(st.progress(0, text="Départ…"))
    for t in range(1, steps + 1):
        for bar, r in zip(bars, rs):
            pct = min(1.0, (_score_from_result(r) / max_score) * (t / steps))
            bar.progress(int(pct * 100), text=f"{int(pct*100)}%")
        time.sleep(sleep)

def _table_results(results: List[Dict[str, Any]], title: str = "📊 Comparatif des réponses") -> None:
    if not results:
        return

    from html import escape as _esc

    ordered = sorted(results, key=_score_from_result, reverse=True)
    rows_all = []
    for i, r in enumerate(ordered, start=1):
        rows_all.append({
            "Rang": i,
            "Modèle": r.get("model", "—"),
            "Score": round(_score_from_result(r), 3),
            "Confiance": round(_safe_num(r.get("confidence", r.get("score", 0.0))), 3),
            "Grounding": round(_safe_num(r.get("grounding", 0.0)), 3),
            "Couverture": round(_safe_num(r.get("coverage", 0.0)), 3),
            "Style": round(_safe_num(r.get("style", 0.0)), 3),
            "Réponse": str(r.get("answer", "")).strip(),
        })

    st.subheader(title)
    st.markdown("""
<style>
.tbl-wrap { width:100%; overflow-x:auto; }
.tbl, .tbl2 {
  width:100%; max-width:1500px; margin:0 auto 18px auto;
  border-collapse: collapse; table-layout: fixed; background:#fff;
  border:1px solid #d1d5db; border-radius:10px; overflow:hidden;
  box-shadow:0 3px 12px rgba(0,0,0,.06);
}
.tbl th, .tbl td, .tbl2 th, .tbl2 td {
  border:1px solid #d1d5db; padding:10px 12px;
  font-size:.98rem; line-height:1.45; text-align:center; vertical-align:middle;
}
.tbl thead th, .tbl2 thead th {
  position: sticky; top: 0; z-index: 1; background:#f1f5f9; color:#0f172a; font-weight:700;
}
.tbl tbody tr:nth-child(even), .tbl2 tbody tr:nth-child(even){ background:#f9fafb; }
.tbl tbody tr:hover, .tbl2 tbody tr:hover{ background:#e0f2fe; transition: background .15s ease; }
.tbl td:nth-child(1), .tbl td:nth-child(2), .tbl td:nth-child(3),
.tbl td:nth-child(4), .tbl td:nth-child(5), .tbl td:nth-child(6), .tbl td:nth-child(7) { white-space:nowrap; }
.tbl th:nth-child(1), .tbl td:nth-child(1) { min-width:80px; }
.tbl th:nth-child(2), .tbl td:nth-child(2) { min-width:220px; }
.tbl th:nth-child(3), .tbl td:nth-child(3) { min-width:110px; }
.tbl th:nth-child(4), .tbl td:nth-child(4) { min-width:130px; }
.tbl th:nth-child(5), .tbl td:nth-child(5) { min-width:130px; }
.tbl th:nth-child(6), .tbl td:nth-child(6) { min-width:130px; }
.tbl th:nth-child(7), .tbl td:nth-child(7) { min-width:110px; }
.tbl2 th:nth-child(1), .tbl2 td:nth-child(1) { min-width:80px; text-align:center; }
.tbl2 th:nth-child(2), .tbl2 td:nth-child(2) { min-width:220px; text-align:center; }
.tbl2 th:nth-child(3), .tbl2 td:nth-child(3) {
  min-width:760px; width:65vw; text-align:left; white-space: normal; word-break: break-word; overflow-wrap: anywhere; line-height: 1.55;
}
</style>
""", unsafe_allow_html=True)

    headers_metrics = ["Rang", "Modèle", "Score", "Confiance", "Grounding", "Couverture", "Style"]
    colgroup_metrics = """
<colgroup>
  <col style="width:90px">
  <col style="width:240px">
  <col style="width:120px">
  <col style="width:140px">
  <col style="width:140px">
  <col style="width:140px">
  <col style="width:120px">
</colgroup>
"""
    html1 = ['<div class="tbl-wrap"><table class="tbl">', colgroup_metrics, "<thead><tr>"]
    for h in headers_metrics:
        html1.append(f"<th>{_esc(h)}</th>")
    html1.append("</tr></thead><tbody>")

    for r in rows_all:
        html1.append("<tr>")
        for h in headers_metrics:
            html1.append(f"<td>{_esc(str(r[h]))}</td>")
        html1.append("</tr>")
    html1.append("</tbody></table></div>")
    st.markdown("".join(html1), unsafe_allow_html=True)

    st.markdown("<div style='height:14px'></div>", unsafe_allow_html=True)

    st.subheader("📝 Réponses détaillées")
    headers_resp = ["Rang", "Modèle", "Réponse"]
    colgroup_resp = """
<colgroup>
  <col style="width:90px">
  <col style="width:240px">
  <col style="width:auto">
</colgroup>
"""
    html2 = ['<div class="tbl-wrap"><table class="tbl2">', colgroup_resp, "<thead><tr>"]
    for h in headers_resp:
        html2.append(f"<th>{_esc(h)}</th>")
    html2.append("</tr></thead><tbody>")

    for r in rows_all:
        html2.append("<tr>")
        html2.append(f"<td>{_esc(str(r['Rang']))}</td>")
        html2.append(f"<td>{_esc(str(r['Modèle']))}</td>")
        html2.append(f"<td>{_esc(str(r['Réponse']))}</td>")
        html2.append("</tr>")

    html2.append("</tbody></table></div>")
    st.markdown("".join(html2), unsafe_allow_html=True)

def _render_summary_quality(report: Dict[str, Any]):
    """
    Affiche les métriques qualité d'un résumé produit par rag_sum.
    'report' contient : tokens_source, tokens_summary, compression, cosine,
                        rouge1_f1, rouge2_f1, keyword_overlap,
                        unsupported, complete_ratio.
    """

    st.subheader("📊 Qualité du résumé")

    tokens_source = report.get("tokens_source", 0)
    tokens_summary = report.get("tokens_summary", 0)
    compression = report.get("compression", 0.0)
    cosine = report.get("cosine", 0.0)
    rouge1 = report.get("rouge1_f1", 0.0)
    rouge2 = report.get("rouge2_f1", 0.0)
    kw = report.get("keyword_overlap", 0.0)
    unsupported = report.get("unsupported", [])
    complete_ratio = report.get("complete_ratio", 0.0)

    col1, col2, col3 = st.columns(3)

    with col1:
        st.metric("📘 Tokens source", tokens_source)
        st.metric("✂️ Compression", f"{compression:.2f}")
        st.metric("🔎 Cosine", f"{cosine:.3f}")

    with col2:
        st.metric("🔤 ROUGE-1 (F1)", f"{rouge1:.3f}")
        st.metric("🔡 ROUGE-2 (F1)", f"{rouge2:.3f}")
        st.metric("🔑 Mots-clés (overlap)", f"{kw:.3f}")

    with col3:
        st.metric("✔️ Phrases complètes", f"{complete_ratio:.2f}")
        st.metric("⚠️ Phrases non ancrées", len(unsupported))

    if unsupported:
        st.warning(
            "Certaines phrases semblent non ancrées dans le texte source :\n\n"
            + "\n".join(f"- {u}" for u in unsupported)
        )


# ============================================================
# Render principal
# ============================================================

def render():
    st.header("🧠 Generate Docs (RAG zéro-hallucination)")

    # Récupérer l'état du dernier résumé (si déjà généré)
    sum_state = st.session_state.get("sum_state")

    mode = st.radio(
        "Mode",
        ["Répondre à une question", "Résumer un document (PDF/TXT)"],
        horizontal=True,
    )

    # Ces variables doivent exister en dehors du form pour le bloc résumé
    up = None
    pasted_text = ""

    with st.form("gen_form"):
        models = _pick_models()
        topk_context = _sidebar_topk_slider(default=6)

        # =====================================================
        # MODE 1 : QUESTION / RAG
        # =====================================================
        if mode == "Répondre à une question":
            q = st.text_area(
                "Question",
                placeholder="Ex: C'est quoi IT-STORM ?",
                height=100,
            )
            c1, c2 = st.columns([1, 1])
            with c1:
                do_gen = st.form_submit_button("Générer (séquentiel)")
            with c2:
                do_cmp = st.form_submit_button("Comparer (jury)")
            do_sum = False
            up = None
            pasted_text = ""
            st.session_state.setdefault("sum_use_template", True)
            st.session_state.setdefault("sum_source_kind", "Choisir…")

        # =====================================================
        # MODE 2 : RÉSUMÉ (DOC OU TEXTE COLLÉ)
        # =====================================================
        else:
            st.session_state.setdefault("sum_use_template", True)
            st.toggle(
                "🧩 Reformuler avec gabarit analytique (global)",
                value=st.session_state["sum_use_template"],
                key="sum_use_template",
                help="Transforme le résumé en un paragraphe fluide (sans inventer), dans un style plus lisible.",
            )

            # 🔥 Rien n’apparaît tant qu’on n’a pas choisi la source
            source_kind = st.radio(
                "Source du contenu à résumer",
                ["Choisir…", "Fichier PDF/TXT", "Texte collé"],
                index=0,
                horizontal=True,
            )
            st.session_state["sum_source_kind"] = source_kind

            if source_kind == "Fichier PDF/TXT":
                up = st.file_uploader(
                    "Importer un document (.pdf ou .txt)",
                    type=["pdf", "txt"],
                )
            elif source_kind == "Texte collé":
                pasted_text = st.text_area(
                    "Texte à résumer",
                    placeholder="Colle ici un ou plusieurs paragraphes (mail, article, texte brut…)",
                    height=220,
                )

            do_sum = st.form_submit_button("Résumer")
            do_gen = do_cmp = False
            q = ""

    # =========================================================
    # TRAITEMENT MODE 1 : QUESTION / RAG
    # =========================================================
    if mode == "Répondre à une question":
        if not (do_gen or do_cmp):
            return

        if not q or not q.strip():
            st.error("Merci de saisir une question.")
            return

        base = _call_smart_rag(q) or {}
        if base:
            _display_grounded_block(base)

        collected: List[Dict[str, Any]] = []
        if do_gen:
            st.subheader("🧪 Réponses par modèle (séquentiel)")
            for m in models:
                with st.spinner(f"{m} génère une réponse."):
                    r = _ask_one_model(q, model=m, topk_context=topk_context, timeout=600.0)
                st.markdown(_decorate_model_label(r.get("model", "—")), unsafe_allow_html=True)
                st.write(r.get("answer", ""))
                collected.append(r)

            if collected:
                best = sorted(collected, key=_score_from_result, reverse=True)[0]
                st.subheader("✅ Réponse finale conseillée")
                st.markdown(_decorate_model_label(best.get("model", "—"), rank=1), unsafe_allow_html=True)
                st.write(best.get("answer", ""))

        results: List[Dict[str, Any]] = []
        if do_cmp:
            if len(models) < 3:
                st.info("Astuce: pour une comparaison significative, sélectionne ≥ 3 modèles.")
            jury = _call_jury(q, models=models, topk_context=topk_context, timeout=600.0) or {}
            results = jury.get("results") or []
            for r in results:
                met = r.get("metrics") or {}
                r["confidence"] = _safe_num(met.get("confidence", r.get("score", 0.0)))
                r["grounding"] = _safe_num(met.get("grounding", 0.0))
                r["coverage"] = _safe_num(met.get("coverage", 0.0))
                r["style"] = _safe_num(met.get("style", 0.0))

        final_list = results or collected
        if final_list:
            _render_podium(final_list)
            _render_race(final_list)
            _table_results(final_list, title="📊 Comparatif des réponses")

        return  # fin du mode question

    # =========================================================
    # TRAITEMENT MODE 2 : RÉSUMÉ (DOC / TEXTE COLLÉ)
    # =========================================================
    if _summarize_text_three_cli is None or read_any_text is None:
        st.error("Module de résumé avancé (rag_sum) indisponible. Vérifie src/rag_sum.py.")
        return

    use_template = bool(st.session_state.get("sum_use_template", True))
    source_kind = st.session_state.get("sum_source_kind", "Choisir…")

    if do_sum:
        # ----- Nouveau calcul de résumé -----
        # 1) Récupération du texte selon la source choisie
        if source_kind == "Fichier PDF/TXT":
            if up is None:
                st.error("Importe un fichier .pdf ou .txt à résumer.")
                return

            file_bytes = up.read()
            filename = up.name or "document"
            text, ext = read_any_text(file_bytes, filename)

            if not text or not text.strip():
                st.error(f"Impossible de lire le contenu de « {filename} ».")
                return

        elif source_kind == "Texte collé":
            raw = (pasted_text or "").strip()
            if not raw:
                st.error("Merci de coller un texte à résumer.")
                return
            text = raw
            ext = "(texte)"
            filename = "Texte collé"

        else:
            st.error("Merci de choisir la source du contenu à résumer.")
            return

        text_len = len(text or "")
        st.info(f"Longueur texte : {text_len} caractères")

        models_for_sum = models or ["mistral:7b-instruct", "llama3.2:3b", "qwen2.5:7b"]

        with st.spinner("⏳ Génération des résumés (mode three-direct, 1 par modèle)…"):
            summaries = _summarize_text_three_cli(text, models=models_for_sum, timeout=200.0) or []

        if not summaries:
            st.error("Aucun résumé n'a été produit (summaries vide).")
            return

        # Meilleur résumé interne rag_sum
        best = max(summaries, key=lambda x: x.get("score", 0.0))
        best_report = best.get("report", {}) or {}
        raw_best_text = best.get("answer", "") or ""
        best_paragraph = _apply_coherence(raw_best_text.strip(), use_template)
        best_export_text = best_paragraph if best_paragraph else raw_best_text.strip()
        best_model_name = best.get("model", "—")

        # Ultimate Version (fusion multi-modèles)
        try:
            ultimate_text = _build_ultimate_summary(
                source_text=text,
                candidates=summaries,
                use_template=use_template,
            )
        except Exception as e:
            st.warning(
                f"Fusion ultime impossible, utilisation de la meilleure version uniquement. Détail : {e}"
            )
            ultimate_text = best_export_text

        # Stockage en session
        st.session_state["sum_state"] = {
            "filename": filename,
            "ext": ext,
            "text_len": text_len,
            "best_model_name": best_model_name,
            "best_export_text": best_export_text,
            "best_report": best_report,
            "use_template": use_template,
            "summaries": summaries,
            "ultimate_text": ultimate_text,
            "source_kind": source_kind,
        }
        sum_state = st.session_state["sum_state"]

    else:
        # Aucun nouveau "Résumer" cliqué : on réutilise ce qui est déjà en mémoire
        if not sum_state:
            st.info("Choisis la source (fichier ou texte collé), puis clique sur « Résumer » pour générer un résumé.")
            return

    # ----- Affichage à partir de sum_state -----
    filename = sum_state.get("filename", "document")
    ext = sum_state.get("ext", "")
    text_len = sum_state.get("text_len")
    best_model_name = sum_state.get("best_model_name", "—")
    best_export_text = sum_state.get("best_export_text", "")
    best_report = sum_state.get("best_report", {}) or {}
    use_template = bool(sum_state.get("use_template", True))
    summaries = sum_state.get("summaries", []) or []
    ultimate_text = sum_state.get("ultimate_text", best_export_text)

    if text_len:
        st.info(f"Longueur texte : {text_len} caractères")

    # ----- Résumé principal -----
    st.subheader("📝 Résumé — meilleure proposition")
    st.markdown(f"**Source :** {filename} ({ext}) · **Modèle sélectionné :** {best_model_name}")
    st.write(best_export_text)

    # ----- Qualité du résumé -----
    if best_report:
        _render_summary_quality(best_report)

    # ----- Ultimate Version -----
    if ultimate_text and ultimate_text != best_export_text:
        st.subheader("✨ Version fusionnée (Ultimate)")
        st.write(ultimate_text)

    # =====================================================
    # Export DOCX : Best Version & Ultimate Version
    # =====================================================
    col_best, col_ultimate = st.columns(2)
    with col_best:
        if best_export_text:
            buf_best = _make_docx_from_summary(
                title="Résumé - Best Version",
                intro=f"Document : {filename} · Modèle sélectionné : {best_model_name}",
                body=best_export_text,
            )
            st.download_button(
                "📄 Exporter — Best Version",
                data=buf_best,
                file_name="resume_best_version.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

    with col_ultimate:
        if ultimate_text:
            buf_ult = _make_docx_from_summary(
                title="Résumé - Ultimate Version",
                intro=(
                    f"Document : {filename} · Version fusionnée à partir des meilleurs résumés "
                    f"(mistral / llama / qwen, etc.)."
                ),
                body=ultimate_text,
            )
            st.download_button(
                "✨ Exporter — Ultimate Version",
                data=buf_ult,
                file_name="resume_ultimate_version.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

    # ----- Tous les résumés par modèle -----
    if summaries:
        st.subheader("📑 Résumés par modèle")
        for i, r in enumerate(sorted(summaries, key=lambda x: x.get("score", 0.0), reverse=True), start=1):
            score = r.get("score", 0.0)
            label = f"{i}. {_label(r.get('model','—'))} ({score:.3f})"
            with st.expander(label, expanded=(i == 1)):
                txt = r.get("answer") or ""
                para = _apply_coherence(txt.strip(), use_template)
                st.write(para if para else txt)

    st.markdown("---")
    if st.button("🔄 Réinitialiser le résumé"):
        st.session_state["sum_state"] = None
        st.rerun()


# Compat app.py
def render_generate_docs_tab():
    render()
